"""
Word report generator for MBTI test results.
Single-page layout — evenly fills an A4 page.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


BLUE = RGBColor(0x2c, 0x5f, 0x8a)
GRAY = RGBColor(0x88, 0x88, 0x88)
ACCENT = RGBColor(0x3c, 0xa4, 0xa4)
LIGHT_GRAY = RGBColor(0xaa, 0xaa, 0xaa)


def _p(doc, text, bold=False, italic=False, size=11, color=None, align=None, before=0, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    return p


def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = "Calibri"
    r.font.color.rgb = BLUE
    return p


def _header_cell(cell, text, fill_color):
    """Style a table header cell with background color and white text."""
    cell.width = Cm(8.25)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): fill_color,
        qn("w:val"): "clear",
        qn("w:color"): "auto",
    })
    tcPr.insert(0, shd)


def _body_cell(cell, text):
    """Style a table body cell."""
    cell.width = Cm(8.25)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if text:
        r = p.add_run("  " + text)
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"


def generate_report(output_path, mbti_type, scores, personality, name=""):
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──
    _p(doc, "MBTI Personality Report", bold=True, size=18, color=BLUE,
       align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    _p(doc, "Myers-Briggs Type Indicator  —  Personal Assessment",
       italic=True, size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    # ── MBTI Type ──
    _p(doc, mbti_type, bold=True, size=32, color=BLUE,
       align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    _p(doc, personality.get("title", ""), italic=True, size=13, color=GRAY,
       align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

    # ── Dimension Breakdown ──
    _h2(doc, "Dimension Breakdown")

    dim_data = [
        ("EI", "E", "I",
         "Extraversion  vs  Introversion",
         "Where you focus your attention and get energy"),
        ("SN", "S", "N",
         "Sensing  vs  Intuition",
         "How you take in information and perceive the world"),
        ("TF", "T", "F",
         "Thinking  vs  Feeling",
         "How you make decisions and evaluate information"),
        ("JP", "J", "P",
         "Judging  vs  Perceiving",
         "How you approach structure and the outer world"),
    ]

    tbl = doc.add_table(rows=len(dim_data), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"

    for i, (dim, pole_a, pole_b, desc, meaning) in enumerate(dim_data):
        pct_a = scores[dim][pole_a]
        pct_b = scores[dim][pole_b]
        bar_len = 22
        bar_a = round(pct_a / 100 * bar_len)
        bar_b = bar_len - bar_a

        # Col 0: dimension name
        c0 = tbl.rows[i].cells[0]
        c0.width = Cm(3.2)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(desc)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.name = "Calibri"

        # Col 1: bar + percentages
        c1 = tbl.rows[i].cells[1]
        c1.width = Cm(8)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)

        r_a = p1.add_run(f"{pole_a}  {'█' * bar_a}  {pct_a}%")
        r_a.bold = True
        r_a.font.size = Pt(9)
        r_a.font.name = "Calibri"
        r_a.font.color.rgb = BLUE

        r_b = p1.add_run(f"    {pole_b}  {'█' * bar_b}  {pct_b}%")
        r_b.bold = True
        r_b.font.size = Pt(9)
        r_b.font.name = "Calibri"
        r_b.font.color.rgb = ACCENT

        # Col 2: meaning
        c2 = tbl.rows[i].cells[2]
        c2.width = Cm(4.8)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(meaning)
        r2.italic = True
        r2.font.size = Pt(8)
        r2.font.name = "Calibri"
        r2.font.color.rgb = GRAY

    # ── Overview ──
    _h2(doc, "Overview")
    _p(doc, personality.get("overview", ""), size=10, after=1)

    # ── Key Traits ──
    _h2(doc, "Key Traits")
    _p(doc, personality.get("traits", ""), size=10, after=1)

    # ── Work Style ──
    _h2(doc, "Work Style & Career")
    _p(doc, personality.get("work_style", ""), size=10, after=1)

    # ── Strengths & Weaknesses ──
    _h2(doc, "Strengths & Areas for Growth")

    strengths = [s.strip().rstrip(".") for s in personality.get("strengths", "").split(",") if s.strip()]
    weaknesses = [w.strip().rstrip(".") for w in personality.get("weaknesses", "").split(",") if w.strip()]
    max_rows = max(len(strengths), len(weaknesses))

    sw_tbl = doc.add_table(rows=max_rows + 1, cols=2)
    sw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sw_tbl.style = "Light Grid Accent 1"

    _header_cell(sw_tbl.rows[0].cells[0], "Strengths", "2c5f8a")
    _header_cell(sw_tbl.rows[0].cells[1], "Areas for Growth", "3ca4a4")

    for i in range(max_rows):
        _body_cell(sw_tbl.rows[i + 1].cells[0], strengths[i] if i < len(strengths) else "")
        _body_cell(sw_tbl.rows[i + 1].cells[1], weaknesses[i] if i < len(weaknesses) else "")

    # ── Footer ──
    _p(doc, "", before=6, after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "This report is generated for informational and educational purposes. "
        "MBTI is a personality indicator, not a diagnostic tool."
    )
    r.italic = True
    r.font.size = Pt(7.5)
    r.font.name = "Calibri"
    r.font.color.rgb = LIGHT_GRAY

    doc.save(output_path)
    print(f"Report saved: {output_path}")
    return output_path
